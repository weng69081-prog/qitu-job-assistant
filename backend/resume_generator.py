"""按模板布局生成 .docx 简历文件"""

import base64
import io
import os
import tempfile
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from typing import Optional


# ── 颜色 ──
CLR_ACCENT = {
    "classic": RGBColor(0x66, 0x7E, 0xEA),
    "skill-first": RGBColor(0xF5, 0x9E, 0x0B),
    "project-focused": RGBColor(0x10, 0xB9, 0x81),
    "fresh-grad": RGBColor(0xEF, 0x44, 0x44),
    "management": RGBColor(0x8B, 0x5C, 0xF6),
    "minimal": RGBColor(0x06, 0xB6, 0xD4),
}
CLR_DARK = RGBColor(0x33, 0x33, 0x33)
CLR_GRAY = RGBColor(0x88, 0x88, 0x88)
CLR_LIGHT = RGBColor(0xF5, 0xF5, 0xF5)
CLR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT_HAN = "微软雅黑"
FONT_EN = "Calibri"


def _set_font(run, name=FONT_HAN, size=10.5, bold=False, color=CLR_DARK, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.italic = italic
    # 正文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), name)


def _add_bottom_border(paragraph, color="667eea", sz=6):
    """给段落加底部边框（用作分割线）"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="4" w:color="{color}"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


def _add_paragraph(doc, text, size=10.5, bold=False, color=None, align=None, spacing_after=6, font_name=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run, name=font_name or FONT_HAN, size=size, bold=bold, color=color or CLR_DARK)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(size * 1.4)
    return p


def _add_section_heading(doc, text, template_id="classic"):
    """不同模板的节标题样式不同"""
    accent = CLR_ACCENT.get(template_id, CLR_ACCENT["classic"])
    if template_id == "minimal":
        p = _add_paragraph(doc, text, size=11, bold=True, color=accent, spacing_after=4)
        p.paragraph_format.space_before = Pt(6)
        # 细线
        _add_bottom_border(p, color="06b6d4", sz=4)
    elif template_id == "project-focused":
        p = _add_paragraph(doc, f"◆ {text} ◆", size=11, bold=True, color=accent, spacing_after=4)
        p.paragraph_format.space_before = Pt(8)
        _add_bottom_border(p, color="10b981", sz=6)
    elif template_id == "fresh-grad":
        p = _add_paragraph(doc, text, size=12, bold=True, color=accent, spacing_after=4)
        p.paragraph_format.space_before = Pt(8)
        _add_bottom_border(p, color="ef4444", sz=8)
    elif template_id == "management":
        p = _add_paragraph(doc, f"▸ {text}", size=11, bold=True, color=accent, spacing_after=4)
        p.paragraph_format.space_before = Pt(8)
        _add_bottom_border(p, color="8b5cf6", sz=6)
    elif template_id == "skill-first":
        p = _add_paragraph(doc, f"━ {text} ━", size=11, bold=True, color=accent, spacing_after=4)
        p.paragraph_format.space_before = Pt(6)
    else:
        p = _add_paragraph(doc, text, size=11, bold=True, color=accent, spacing_after=4)
        p.paragraph_format.space_before = Pt(6)
        _add_bottom_border(p, color="667eea", sz=6)


def _add_info_line(doc, label, value, template_id="classic", indent=0):
    """添加一行信息，不同模板不同前缀"""
    if not value:
        return
    prefix = " " * indent
    if template_id == "management":
        text = f"▸ {label}：{value}"
    elif template_id == "skill-first":
        text = f"• {label}：{value}"
    elif template_id == "project-focused":
        text = f"  {label}：{value}"
    elif template_id == "minimal":
        text = f"  {label}：{value}"
    else:
        text = f"- {label}：{value}"
    _add_paragraph(doc, text, size=10)


def _add_multi_items(doc, items, template_id, name_field="name", period_field="period",
                     content_field="content", content_label="", role_field=""):
    """添加多条条目（工作经历、项目经历等）"""
    accent = CLR_ACCENT.get(template_id, CLR_ACCENT["classic"])
    for item in items:
        name = getattr(item, name_field, "") or ""
        period = getattr(item, period_field, "") or ""
        content = getattr(item, content_field, "") or ""
        role = getattr(item, role_field, "") if role_field else ""

        # 标题行
        title = name
        if period:
            if template_id == "project-focused":
                title = f"■■ {name}  ══ {period} ══"
            elif template_id == "management":
                title = f"▸ {name}  📅 {period}"
            elif template_id == "minimal":
                title = f"  {name}  ·  {period}"
            else:
                title = f"{name}  ({period})"

        t = _add_paragraph(doc, title, size=10.5, bold=True, spacing_after=2)

        if role:
            if template_id == "project-focused":
                _add_paragraph(doc, f"  → 职责：{role}", size=9.5, color=CLR_GRAY, spacing_after=1)
            else:
                _add_paragraph(doc, f"  {role}", size=9.5, color=CLR_GRAY, spacing_after=1)

        if content:
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    if template_id == "project-focused":
                        _add_paragraph(doc, f"  → {line}", size=10, spacing_after=1)
                    elif template_id == "management":
                        _add_paragraph(doc, f"    │ {line}", size=10, spacing_after=1)
                    elif template_id == "minimal":
                        _add_paragraph(doc, f"    {line}", size=10, spacing_after=1)
                    else:
                        _add_paragraph(doc, f"  · {line}", size=10, spacing_after=1)

        _add_paragraph(doc, "", size=6, spacing_after=2)  # 间距


def _remove_table_borders(table):
    """移除表格所有边框"""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_photo_to_cell(cell, photo_bytes):
    """在表格单元格中添加照片（居中）"""
    with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False, dir="/tmp") as tmp:
        tmp.write(photo_bytes)
        tmp_path = tmp.name
    try:
        run = cell.paragraphs[0].add_run()
        run.add_picture(tmp_path, width=Cm(2.5), height=Cm(3.5))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 垂直居中
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
            tc.insert(0, tcPr)
        vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
        tcPr.append(vAlign)
    finally:
        os.unlink(tmp_path)


def _add_photo_header(doc, name_text, career, requirements, accent, tid, photo_bytes):
    """创建含照片的顶部表格：左侧姓名+求职意向，右侧照片"""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    left_cell.width = Cm(14.0)
    right_cell.width = Cm(3.5)

    _remove_table_borders(table)

    # 左列：姓名
    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name_text)
    _set_font(run, size=18, bold=True, color=accent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

    # 左列：求职意向
    if career:
        p2 = left_cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"求职意向：{career}")
        _set_font(run2, size=10, color=CLR_GRAY)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.space_before = Pt(0)

    # 左列：需求
    if requirements:
        p3 = left_cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(f"需求：{requirements}")
        _set_font(run3, size=9, color=CLR_GRAY)
        p3.paragraph_format.space_after = Pt(4)

    # 简约现代装饰线（放在左列底部）
    if tid == "minimal":
        p_deco = left_cell.add_paragraph()
        p_deco.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_d = p_deco.add_run("· · · · · · · · · · · · · · · · · · · · ·")
        _set_font(run_d, size=8, color=CLR_GRAY)
        p_deco.paragraph_format.space_after = Pt(8)

    # 右列：照片
    _add_photo_to_cell(right_cell, photo_bytes)

    # 表格后加一小段间距
    _add_paragraph(doc, "", size=2, spacing_after=2)


def generate_resume_docx(data, output_path: str, photo_bytes: Optional[bytes] = None) -> str:
    """生成简历 .docx 文件，返回文件路径"""
    tid = data.template_id or "classic"
    accent = CLR_ACCENT.get(tid, CLR_ACCENT["classic"])

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    if tid == "minimal":
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    elif tid == "management":
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    else:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── 标题（姓名）+ 照片（有则用表格，无则纯文字） ──
    name_text = data.name or "个人简历"
    career = data.career or data.job_intention or ""

    if photo_bytes:
        # 含照片：新建表格布局（左文字右照片）
        _add_photo_header(doc, name_text, career, data.requirements, accent, tid, photo_bytes)
    else:
        # 无照片：纯文字
        t = _add_paragraph(doc, name_text, size=18, bold=True, color=accent,
                           align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
        if career:
            _add_paragraph(doc, f"求职意向：{career}", size=10, color=CLR_GRAY,
                           align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
        if data.requirements:
            _add_paragraph(doc, f"需求：{data.requirements}", size=9, color=CLR_GRAY,
                           align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=4)
        if tid == "minimal":
            _add_paragraph(doc, "· · · · · · · · · · · · · · · · · · · · ·", size=8, color=CLR_GRAY,
                           align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=8)

    # ══════════════════════════════════════════
    #  1. 个人信息
    # ══════════════════════════════════════════

    if tid == "skill-first":
        # 两栏：左列个人信息+技能，右列经历
        _render_skill_first(doc, data, tid, career, photo_bytes)
    else:
        _add_section_heading(doc, "个人信息", tid)

        info_items = []
        if data.name: info_items.append(("姓名", data.name))
        if data.gender: info_items.append(("性别", data.gender))
        if data.age: info_items.append(("年龄", data.age))
        if data.phone: info_items.append(("电话", data.phone))
        if data.email: info_items.append(("邮箱", data.email))
        if data.address: info_items.append(("现居", data.address))
        if data.wechat: info_items.append(("微信", data.wechat))

        # 信息用一行显示（紧凑模式）
        if tid in ("management", "minimal"):
            parts = [f"{l}：{v}" for l, v in info_items]
            _add_paragraph(doc, "  |  ".join(parts), size=10, spacing_after=6)
        else:
            for label, value in info_items:
                _add_info_line(doc, label, value, tid)

        if career and not data.career and not data.job_intention:
            pass  # 已经在上方显示了
        _add_paragraph(doc, "", size=4, spacing_after=2)

        # ══════════════════════════════════════════
        #  2. 教育经历
        # ══════════════════════════════════════════
        edu = data.education
        if edu.school or edu.major:
            _add_section_heading(doc, "教育经历", tid)
            edu_line = edu.school or ""
            if edu.major: edu_line += f"  |  {edu.major}"
            if edu.degree: edu_line += f"  |  {edu.degree}"
            _add_paragraph(doc, edu_line, size=10.5, bold=True, spacing_after=2)
            if edu.start_date or edu.end_date:
                _add_paragraph(doc, f"{edu.start_date} — {edu.end_date}", size=9.5, color=CLR_GRAY, spacing_after=2)
            if edu.courses:
                _add_paragraph(doc, f"主修课程：{edu.courses}", size=10, spacing_after=4)
            _add_paragraph(doc, "", size=4, spacing_after=2)

        # ══════════════════════════════════════════
        #  3. 工作/实习经历
        # ══════════════════════════════════════════
        if data.experiences:
            _add_section_heading(doc, "工作/实习经历", tid)
            _add_multi_items(doc, data.experiences, tid,
                             name_field="company", period_field="period",
                             content_field="content")

        # ══════════════════════════════════════════
        #  4. 项目经历
        # ══════════════════════════════════════════
        if data.projects:
            _add_section_heading(doc, "项目经历", tid)
            _add_multi_items(doc, data.projects, tid,
                             name_field="name", period_field="period",
                             content_field="achievement", role_field="role")

        # ══════════════════════════════════════════
        #  5. 专业技能
        # ══════════════════════════════════════════
        sk = data.skills
        has_skill = any([sk.office, sk.professional, sk.programming, sk.certificates, sk.languages])
        if has_skill:
            _add_section_heading(doc, "专业技能", tid)
            skill_parts = []
            if sk.office: skill_parts.append(f"办公软件：{sk.office}")
            if sk.professional: skill_parts.append(f"专业软件：{sk.professional}")
            if sk.programming: skill_parts.append(f"编程语言：{sk.programming}")
            if sk.certificates: skill_parts.append(f"证书：{sk.certificates}")
            if sk.languages: skill_parts.append(f"语言能力：{sk.languages}")
            for s in skill_parts:
                if tid in ("management", "minimal"):
                    _add_paragraph(doc, f"  {s}", size=10, spacing_after=1)
                else:
                    _add_paragraph(doc, f"- {s}", size=10, spacing_after=1)
            _add_paragraph(doc, "", size=4, spacing_after=2)

        # ══════════════════════════════════════════
        #  6. 校园经历
        # ══════════════════════════════════════════
        if data.campus_activities:
            _add_section_heading(doc, "校园经历", tid)
            for act in data.campus_activities:
                header = act.name or ""
                if act.period: header += f"  ·  {act.period}"
                _add_paragraph(doc, header, size=10.5, bold=True, spacing_after=1)
                if act.role:
                    _add_paragraph(doc, f"  {act.role}", size=10, spacing_after=2)
            _add_paragraph(doc, "", size=4, spacing_after=2)

        # ══════════════════════════════════════════
        #  7. 荣誉奖项
        # ══════════════════════════════════════════
        if data.awards:
            _add_section_heading(doc, "荣誉奖项", tid)
            for aw in data.awards:
                line = aw.name or ""
                extra = ""
                if aw.period: extra = f"（{aw.period}）"
                if aw.issuer: extra += f"  颁发：{aw.issuer}"
                if tid == "management":
                    _add_paragraph(doc, f"🏆 {line} {extra}".strip(), size=10, spacing_after=1)
                elif tid == "minimal":
                    _add_paragraph(doc, f"  {line} {extra}".strip(), size=10, spacing_after=1)
                else:
                    _add_paragraph(doc, f"- {line} {extra}".strip(), size=10, spacing_after=1)
            _add_paragraph(doc, "", size=4, spacing_after=2)

        # ══════════════════════════════════════════
        #  8. 自我评价
        # ══════════════════════════════════════════
        if data.self_evaluation:
            _add_section_heading(doc, "自我评价", tid)
            for line in data.self_evaluation.split("\n"):
                line = line.strip()
                if line:
                    if tid == "management":
                        _add_paragraph(doc, f"💬 {line}", size=10, spacing_after=1)
                    else:
                        _add_paragraph(doc, f"{line}", size=10, spacing_after=1)

    # ── 保存 ──
    doc.save(output_path)
    return output_path


def _render_skill_first(doc, data, tid, career, photo_bytes=None):
    """技能突出模板：两栏布局"""

    # 标题
    name_text = data.name or "个人简历"
    t = _add_paragraph(doc, name_text, size=20, bold=True,
                       color=CLR_ACCENT["skill-first"],
                       align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=2)
    if career:
        _add_paragraph(doc, f"求职意向：{career}", size=10, color=CLR_GRAY,
                       align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=8)

    # 两栏表格 (30% / 70%)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 设置列宽
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # 左栏宽度
    left_cell.width = Cm(5.0)
    right_cell.width = Cm(12.0)

    # 左栏背景色（浅黄）
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FEF7E6" w:val="clear"/>')
    left_cell._element.get_or_add_tcPr().append(shading)

    # 照片（放在左栏顶部）
    if photo_bytes:
        _add_photo_to_cell(left_cell, photo_bytes)
        _add_paragraph_in_cell(left_cell, "", size=4, spacing_after=2)

    # 左栏：个人信息
    p = _add_paragraph_in_cell(left_cell, "个人信息", size=11, bold=True,
                                color=CLR_ACCENT["skill-first"], spacing_after=4)
    _add_bottom_border(p, color="f59e0b", sz=4)

    for label, value in [("姓名", data.name), ("性别", data.gender), ("年龄", data.age),
                          ("电话", data.phone), ("邮箱", data.email), ("现居", data.address),
                          ("微信", data.wechat)]:
        if value:
            _add_paragraph_in_cell(left_cell, f"• {label}：{value}", size=9, spacing_after=1)

    # 左栏：专业技能（大区块）
    sk = data.skills
    if any([sk.office, sk.professional, sk.programming, sk.certificates, sk.languages]):
        _add_paragraph_in_cell(left_cell, "", size=6, spacing_after=2)
        p = _add_paragraph_in_cell(left_cell, "专业技能", size=11, bold=True,
                                    color=CLR_ACCENT["skill-first"], spacing_after=4)
        _add_bottom_border(p, color="f59e0b", sz=4)

        skill_lines = []
        if sk.office: skill_lines.append(f"办公：{sk.office}")
        if sk.professional: skill_lines.append(f"专业：{sk.professional}")
        if sk.programming: skill_lines.append(f"编程：{sk.programming}")
        if sk.certificates: skill_lines.append(f"证书：{sk.certificates}")
        if sk.languages: skill_lines.append(f"语言：{sk.languages}")
        for s in skill_lines:
            _add_paragraph_in_cell(left_cell, f"• {s}", size=9, spacing_after=2)

    # 右栏：教育经历
    edu = data.education
    if edu.school or edu.major:
        p = _add_paragraph_in_cell(right_cell, "教育经历", size=11, bold=True,
                                    color=CLR_ACCENT["skill-first"], spacing_after=2)
        edu_line = edu.school or ""
        if edu.major: edu_line += f"  |  {edu.major}"
        if edu.degree: edu_line += f"  |  {edu.degree}"
        _add_paragraph_in_cell(right_cell, edu_line, size=10, bold=True, spacing_after=1)
        if edu.start_date or edu.end_date:
            _add_paragraph_in_cell(right_cell, f"{edu.start_date} — {edu.end_date}", size=9, color=CLR_GRAY, spacing_after=1)
        if edu.courses:
            _add_paragraph_in_cell(right_cell, f"课程：{edu.courses}", size=9, spacing_after=4)
        _add_paragraph_in_cell(right_cell, "", size=2, spacing_after=1)

    # 右栏：工作经历
    if data.experiences:
        p = _add_paragraph_in_cell(right_cell, "工作经历", size=11, bold=True,
                                    color=CLR_ACCENT["skill-first"], spacing_after=2)
        _add_multi_items_in_cell(right_cell, data.experiences, tid="skill-first",
                                 name_field="company", period_field="period",
                                 content_field="content")

    # 右栏：项目经历
    if data.projects:
        _add_paragraph_in_cell(right_cell, "", size=2, spacing_after=1)
        p = _add_paragraph_in_cell(right_cell, "项目经历", size=11, bold=True,
                                    color=CLR_ACCENT["skill-first"], spacing_after=2)
        _add_multi_items_in_cell(right_cell, data.projects, tid="skill-first",
                                 name_field="name", period_field="period",
                                 content_field="achievement", role_field="role")

    # 右栏：校园经历（如果有）
    if data.campus_activities:
        _add_paragraph_in_cell(right_cell, "", size=2, spacing_after=1)
        p = _add_paragraph_in_cell(right_cell, "校园经历", size=11, bold=True,
                                    color=CLR_ACCENT["skill-first"], spacing_after=2)
        for act in data.campus_activities:
            header = act.name or ""
            if act.period: header += f"  ·  {act.period}"
            _add_paragraph_in_cell(right_cell, header, size=10, bold=True, spacing_after=1)
            if act.role:
                _add_paragraph_in_cell(right_cell, f"  {act.role}", size=9, spacing_after=2)

    # 移除表格边框
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_paragraph_in_cell(cell, text, size=10, bold=False, color=None, align=None, spacing_after=4):
    """在表格cell里加段落"""
    p = cell.add_paragraph()
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color or CLR_DARK)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(size * 1.35)
    return p


def _add_multi_items_in_cell(cell, items, tid, name_field="name", period_field="period",
                              content_field="content", role_field=""):
    """在cell里添加条目"""
    for item in items:
        name = getattr(item, name_field, "") or ""
        period = getattr(item, period_field, "") or ""
        content = getattr(item, content_field, "") or ""
        role = getattr(item, role_field, "") if role_field else ""

        title = name
        if period:
            title = f"{name}  ({period})"
        _add_paragraph_in_cell(cell, title, size=10, bold=True, spacing_after=1)

        if role:
            _add_paragraph_in_cell(cell, f"  {role}", size=9, color=CLR_GRAY, spacing_after=1)

        if content:
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    _add_paragraph_in_cell(cell, f"  • {line}", size=9, spacing_after=1)

        _add_paragraph_in_cell(cell, "", size=4, spacing_after=1)